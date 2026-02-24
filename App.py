import streamlit as st
import pandas as pd
import base64
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches

# --- 1. CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="Governança de Dados: Gerador de Stack", layout="wide")
st.title("Arquitetura de dados: Ecossistema e Ferramentas")
st.markdown("Gere a arquitetura de referência ideal, explore a tabela de ferramentas e baixe o relatório executivo completo (.docx).")

# --- 2. BASE DE CONHECIMENTO (As 37 Ferramentas) ---
dados = [
    {"nome": "OpenMetadata", "categoria": "Catálogo / Governança", "tipo": "Open Source", "link": "https://open-metadata.org/", "resumo": "Catálogo API-first, focado em colaboração e linhagem.", "facilidade_instalacao": 5, "suporte": 4, "flexibilidade": 5},
    {"nome": "DataHub", "categoria": "Catálogo / Governança", "tipo": "Open Source", "link": "https://datahubproject.io/", "resumo": "Plataforma robusta (LinkedIn), excelente para ecossistemas complexos.", "facilidade_instalacao": 3, "suporte": 4, "flexibilidade": 5},
    {"nome": "Apache Atlas", "categoria": "Catálogo / Governança", "tipo": "Open Source", "link": "https://atlas.apache.org/", "resumo": "Governança profunda e segurança, nativo do ecossistema Hadoop.", "facilidade_instalacao": 2, "suporte": 3, "flexibilidade": 4},
    {"nome": "Atlan", "categoria": "Catálogo / Governança", "tipo": "Comercial / Pago", "link": "https://atlan.com/", "resumo": "Plataforma SaaS super intuitiva, ideal para times de negócios.", "facilidade_instalacao": 5, "suporte": 5, "flexibilidade": 3},
    {"nome": "Cube", "categoria": "Camada Semântica", "tipo": "Open Source", "link": "https://cube.dev/", "resumo": "Camada semântica ('headless BI') que centraliza métricas via API.", "facilidade_instalacao": 4, "suporte": 4, "flexibilidade": 5},
    {"nome": "dbt Semantic Layer", "categoria": "Camada Semântica", "tipo": "Comercial / Pago", "link": "https://www.getdbt.com/product/semantic-layer", "resumo": "Define métricas diretamente nos modelos de transformação dbt.", "facilidade_instalacao": 4, "suporte": 5, "flexibilidade": 4},
    {"nome": "Looker", "categoria": "Camada Semântica", "tipo": "Comercial / Pago", "link": "https://cloud.google.com/looker", "resumo": "Plataforma do Google com linguagem LookML para modelagem forte.", "facilidade_instalacao": 5, "suporte": 5, "flexibilidade": 4},
    {"nome": "Apache Iceberg", "categoria": "Formato de Tabela", "tipo": "Open Source", "link": "https://iceberg.apache.org/", "resumo": "Formato aberto que traz transações seguras para Data Lakes.", "facilidade_instalacao": 3, "suporte": 4, "flexibilidade": 5},
    {"nome": "Delta Lake", "categoria": "Formato de Tabela", "tipo": "Open Source", "link": "https://delta.io/", "resumo": "Traz confiabilidade e performance ACID ao Lakehouse.", "facilidade_instalacao": 4, "suporte": 5, "flexibilidade": 4},
    {"nome": "Apache Hudi", "categoria": "Formato de Tabela", "tipo": "Open Source", "link": "https://hudi.apache.org/", "resumo": "Otimizado para atualizações constantes (upserts) em tempo real.", "facilidade_instalacao": 3, "suporte": 4, "flexibilidade": 4},
    {"nome": "Parquet", "categoria": "Formato de Arquivo", "tipo": "Open Source", "link": "https://parquet.apache.org/", "resumo": "Armazenamento colunar, hiper-comprimido, padrão para análises.", "facilidade_instalacao": 5, "suporte": 5, "flexibilidade": 3},
    {"nome": "Apache Spark", "categoria": "Processamento / Big Data", "tipo": "Open Source", "link": "https://spark.apache.org/", "resumo": "Motor líder para processar clusters massivos de dados.", "facilidade_instalacao": 2, "suporte": 5, "flexibilidade": 5},
    {"nome": "Trino", "categoria": "Motor de Consulta", "tipo": "Open Source", "link": "https://trino.io/", "resumo": "Motor SQL para nuvem. Consulta dados pesados direto no Lake.", "facilidade_instalacao": 2, "suporte": 5, "flexibilidade": 5},
    {"nome": "DuckDB", "categoria": "Motor de Consulta", "tipo": "Open Source", "link": "https://duckdb.org/", "resumo": "Banco analítico ultrarrápido que roda localmente na máquina do analista.", "facilidade_instalacao": 5, "suporte": 5, "flexibilidade": 5},
    {"nome": "Oracle", "categoria": "Banco de Dados", "tipo": "Comercial / Pago", "link": "https://www.oracle.com/database/", "resumo": "Banco relacional tradicional, altíssima resiliência e custo.", "facilidade_instalacao": 1, "suporte": 5, "flexibilidade": 2},
    {"nome": "Apache Kafka", "categoria": "Streaming de Dados", "tipo": "Open Source", "link": "https://kafka.apache.org/", "resumo": "Plataforma de streaming de eventos em tempo real.", "facilidade_instalacao": 2, "suporte": 5, "flexibilidade": 5},
    {"nome": "Schema Registry", "categoria": "Streaming de Dados", "tipo": "Open Source", "link": "https://docs.confluent.io/platform/current/schema-registry/index.html", "resumo": "Garante que eventos sigam um formato de contrato rígido.", "facilidade_instalacao": 3, "suporte": 5, "flexibilidade": 4},
            # --- NOVA FERRAMENTA ADICIONADA AQUI ---
    {"nome": "Scriptcase", "categoria": "Aplicações e BI (RAD)", "tipo": "Comercial / Pago", "link": "https://www.scriptcase.com.br/", "resumo": "Plataforma de desenvolvimento rápido (RAD) em PHP para criar sistemas web, relatórios e painéis analíticos integrados a bancos de dados.", "facilidade_instalacao": 4, "suporte": 4, "flexibilidade": 4},
    # --- NOVA FERRAMENTA: PENTAHO ---
    {"nome": "Pentaho", "categoria": "Integração e BI", "tipo": "Open Source", "link": "https://pentaho.com/", "resumo": "Plataforma consolidada de Integração de Dados (ETL) e Business Intelligence, que permite construir pipelines complexos através de uma interface visual.", "facilidade_instalacao": 3, "suporte": 4, "flexibilidade": 4},
# --- NOVA FERRAMENTA: POSTGRESQL ---
    {"nome": "PostgreSQL", "categoria": "Banco de Dados", "tipo": "Open Source", "link": "https://www.postgresql.org/", "resumo": "Banco de dados relacional open source mais avançado do mundo, conhecido por sua altíssima confiabilidade, integridade de dados e extensibilidade.", "facilidade_instalacao": 4, "suporte": 5, "flexibilidade": 5},
    # --- NOVA FERRAMENTA: ALATION ---
    {"nome": "Alation", "categoria": "Catálogo / Governança", "tipo": "Comercial / Pago", "link": "https://www.alation.com/", "resumo": "Plataforma líder em inteligência de dados empresarial, com forte foco em catálogo colaborativo, governança ativa e adoção por usuários de negócios.", "facilidade_instalacao": 4, "suporte": 5, "flexibilidade": 4},
    # --- NOVA FERRAMENTA: ATACCAMA ONE ---
    {"nome": "Ataccama ONE", "categoria": "Catálogo / Governança", "tipo": "Comercial / Pago", "link": "https://www.ataccama.com/", "resumo": "Plataforma unificada que integra qualidade de dados, Master Data Management (MDM) e catálogo com automação baseada em IA.", "facilidade_instalacao": 3, "suporte": 5, "flexibilidade": 4},
    # --- NOVA FERRAMENTA: COLLIBRA ---
    {"nome": "Collibra", "categoria": "Catálogo / Governança", "tipo": "Comercial / Pago", "link": "https://www.collibra.com/", "resumo": "Plataforma enterprise de altíssimo nível focada em governança, gestão de políticas, glossário de negócios e fluxos de aprovação (workflows).", "facilidade_instalacao": 3, "suporte": 5, "flexibilidade": 4},
    # --- NOVA FERRAMENTA: ERWIN DATA INTELLIGENCE ---
    {"nome": "Erwin Data Intelligence", "categoria": "Catálogo / Governança", "tipo": "Comercial / Pago", "link": "https://www.quest.com/solutions/data", "resumo": "Plataforma corporativa tradicional que unifica modelagem de dados, catálogo e linhagem automatizada, criando uma ponte sólida entre TI e negócios.", "facilidade_instalacao": 3, "suporte": 5, "flexibilidade": 4},
    # --- NOVA FERRAMENTA: IBM Cloud Pak ---
    {"nome": "IBM Cloud Pak for Data", "categoria": "Plataforma de Dados e IA", "tipo": "Comercial / Pago", "link": "https://www.ibm.com/products/cloud-pak-for-data", "resumo": "Plataforma corporativa unificada de Data Fabric. Centraliza governança avançada e oferece um ambiente completo para a validação e construção de modelos estatísticos e machine learning.", "facilidade_instalacao": 2, "suporte": 5, "flexibilidade": 4},
    # --- NOVA FERRAMENTA: ROCKET DATA INTELLIGENCE ---
    {"nome": "Rocket Data Intelligence", "categoria": "Catálogo / Governança", "tipo": "Comercial / Pago", "link": "https://www.rocketsoftware.com/products/rocket-data-intelligence", "resumo": "Plataforma avançada de descoberta e governança (antiga ASG). Destaca-se por mapear a linhagem de ponta a ponta, conectando sistemas legados aos ecossistemas de nuvem modernos.", "facilidade_instalacao": 3, "suporte": 5, "flexibilidade": 4},
    # --- NOVA FERRAMENTA: OKD KUBERNETES ---
    {"nome": "OKD (Kubernetes)", "categoria": "Infraestrutura / Orquestração", "tipo": "Open Source", "link": "https://www.okd.io/", "resumo": "A distribuição comunitária do Kubernetes que baseia o Red Hat OpenShift. Orquestra containers com alta escalabilidade, ideal para suportar microsserviços e motores de Big Data.", "facilidade_instalacao": 2, "suporte": 4, "flexibilidade": 5},
    # --- NOVA FERRAMENTA: DOCKER ---
    {"nome": "Docker", "categoria": "Infraestrutura / Orquestração", "tipo": "Open Source", "link": "https://www.docker.com/", "resumo": "Plataforma padrão da indústria para a criação e execução de aplicações em contentores isolados. Garante que o ambiente de pesquisa e os modelos estatísticos corram de forma idêntica em qualquer máquina.", "facilidade_instalacao": 4, "suporte": 5, "flexibilidade": 5},
    # --- NOVAS FERRAMENTAS: ESTEIRA DE DEVOPS E REGISTRO ---
    {"nome": "GitLab", "categoria": "DevOps / CI/CD", "tipo": "Open Source", "link": "https://about.gitlab.com/", "resumo": "Plataforma completa de DevOps em uma única aplicação. Oferece repositório de código fonte (Git) robusto, gestão de projetos e esteiras de CI/CD nativas.", "facilidade_instalacao": 3, "suporte": 5, "flexibilidade": 5},
    {"nome": "GitLab Runner", "categoria": "DevOps / CI/CD", "tipo": "Open Source", "link": "https://docs.gitlab.com/runner/", "resumo": "Agente de execução leve e escalável que roda os pipelines de CI/CD definidos no GitLab, automatizando a construção, os testes e o deploy da arquitetura.", "facilidade_instalacao": 4, "suporte": 5, "flexibilidade": 5},
    {"nome": "GitLab Container Registry", "categoria": "DevOps / CI/CD", "tipo": "Open Source", "link": "https://docs.gitlab.com/ee/user/packages/container_registry/", "resumo": "Registro de contêineres seguro e totalmente integrado ao GitLab. Permite armazenar e gerenciar imagens Docker de forma contínua junto ao código-fonte.", "facilidade_instalacao": 4, "suporte": 5, "flexibilidade": 4},
    {"nome": "Harbor", "categoria": "Infraestrutura / Orquestração", "tipo": "Open Source", "link": "https://goharbor.io/", "resumo": "Registro de contêineres cloud-native confiável. Ele armazena, assina e escaneia imagens Docker em busca de vulnerabilidades, sendo o parceiro ideal do Kubernetes para segurança.", "facilidade_instalacao": 3, "suporte": 4, "flexibilidade": 4},
    # --- NOVAS FERRAMENTAS: AJUSTES DO DIAGRAMA ---
    {"nome": "Apache Airflow", "categoria": "Orquestração / Ingestão", "tipo": "Open Source", "link": "https://airflow.apache.org/", "resumo": "Plataforma padrão-ouro para programar, orquestrar e monitorar fluxos de trabalho (pipelines) de dados em lote.", "facilidade_instalacao": 3, "suporte": 5, "flexibilidade": 5},
    {"nome": "Fivetran", "categoria": "Orquestração / Ingestão", "tipo": "Comercial / Pago", "link": "https://www.fivetran.com/", "resumo": "Serviço gerenciado que sincroniza dados de diversas fontes para o seu Data Warehouse com zero configuração (ELT).", "facilidade_instalacao": 5, "suporte": 5, "flexibilidade": 3},
    {"nome": "Snowflake", "categoria": "Processamento / Banco de Dados", "tipo": "Comercial / Pago", "link": "https://www.snowflake.com/", "resumo": "Plataforma de dados em nuvem (Data Cloud) que funciona como Data Warehouse gerenciado, ideal para processamento analítico em lote.", "facilidade_instalacao": 5, "suporte": 5, "flexibilidade": 4},
    # --- NOVAS FERRAMENTAS: BANCOS NOSQL E GRAFOS ---
    {"nome": "MongoDB", "categoria": "Banco de Dados", "tipo": "Open Source", "link": "https://www.mongodb.com/", "resumo": "Banco de dados NoSQL orientado a documentos (JSON). Ideal para receber dados não-estruturados ou semi-estruturados com grande flexibilidade.", "facilidade_instalacao": 4, "suporte": 5, "flexibilidade": 5},
    {"nome": "Neo4j", "categoria": "Banco de Dados", "tipo": "Open Source", "link": "https://neo4j.com/", "resumo": "Banco de dados em grafos líder de mercado. Perfeito para mapear relacionamentos complexos, ligações estruturais e redes de conexão.", "facilidade_instalacao": 3, "suporte": 5, "flexibilidade": 4}
    ]
df = pd.DataFrame(dados)

# --- 3. INTERFACE DE PERGUNTAS (SIDEBAR) ---
st.sidebar.header("Filtros e Diretrizes")

categorias_disponiveis = ["Todas"] + sorted(list(df['categoria'].unique()))
filtro_categoria = st.sidebar.selectbox("0. Filtrar Tabela por Categoria?", categorias_disponiveis)

st.sidebar.markdown("---")
modelo_licenca = st.sidebar.radio("1. Preferência de licenciamento?", ("Mostrar Todos", "Apenas Open Source", "Apenas Comercial / Pago"))
cenario_dados = st.sidebar.radio("2. Comportamento dos dados?", ("Processamento em Lote (Lakehouse, Histórico)", "Streaming em Tempo Real (Eventos)"))

# --- FILTRO NATUREZA DO ARMAZENAMENTO POLIGLOTA ---
natureza_dados = st.sidebar.multiselect(
    "2.1. Como a organização lida com o armazenamento principal? (Pode marcar mais de uma opção)",
    [
        "Dados altamente estruturados (Relacional / Metadados)",
        "Dados não-estruturados ou flexíveis (Documentos, Textos longos, Imagens)",
        "Foco em mapear conexões e relacionamentos complexos (Grafos)",
        "Análise histórica massiva (Data Lakehouse)"
    ],
    default=["Dados altamente estruturados (Relacional / Metadados)"]
)

# --- FILTRO FOCO DA GOVERNANÇA ---
foco_governanca = st.sidebar.radio(
    "Qual é o principal desafio de Governança da organização?",
    (
        "Linhagem técnica e mapeamento para engenharia (Foco em TI)",
        "Glossário, políticas e democratização (Foco em Negócios/Compliance)",
        "Qualidade de dados automatizada e Master Data Management (Foco em MDM)"
    )
)

exige_semantica = st.sidebar.checkbox("3. Implementar Camada Semântica?", value=True)

st.sidebar.markdown("---")
nota_suporte = st.sidebar.slider("Suporte Mínimo (1 a 5)", 1, 5, 1)
nota_flexibilidade = st.sidebar.slider("Flexibilidade Mínima (1 a 5)", 1, 5, 1)

# --- 4. MOTOR DA ARQUITETURA ---
st.subheader("🗺️ Desenho da Arquitetura Recomendada")

camada_ingestao, camada_armazenamento, camada_processamento, camada_semantica_tool, camada_governanca, justificativa = "", "", "", "", "", ""

# 1. Lógica do Armazenamento Poliglota (Banco de Dados)
bancos_selecionados = []
justificativas_banco = []

if "Dados altamente estruturados (Relacional / Metadados)" in natureza_dados:
    banco_rel = "Oracle" if modelo_licenca == "Apenas Comercial / Pago" else "PostgreSQL"
    bancos_selecionados.append(banco_rel)
    justificativas_banco.append(f"O {banco_rel} garantirá a integridade relacional e transacional rígida dos metadados. ")

if "Dados não-estruturados ou flexíveis (Documentos, Textos longos, Imagens)" in natureza_dados:
    bancos_selecionados.append("MongoDB")
    justificativas_banco.append("O MongoDB trará a flexibilidade NoSQL necessária para acomodar documentos variáveis e indexação de arquivos. ")

if "Foco em mapear conexões e relacionamentos complexos (Grafos)" in natureza_dados:
    bancos_selecionados.append("Neo4j")
    justificativas_banco.append("O Neo4j foi incluído para permitir consultas de altíssima performance em vínculos e relacionamentos estruturais complexos. ")

if "Análise histórica massiva (Data Lakehouse)" in natureza_dados:
    banco_lake = "Snowflake" if modelo_licenca == "Apenas Comercial / Pago" else "Apache Iceberg"
    bancos_selecionados.append(banco_lake)
    justificativas_banco.append(f"O {banco_lake} servirá como repositório analítico escalável para o cruzamento de dados históricos em massa. ")

# Prevenção: Se o usuário desmarcar tudo, define um padrão
if not bancos_selecionados:
    bancos_selecionados = ["Apache Iceberg"]
    justificativas_banco = ["O Apache Iceberg foi definido como repositório padrão. "]

# Junta as ferramentas com o sinal de soma (+) para o diagrama visual
camada_armazenamento = " + ".join(bancos_selecionados)
justificativa_banco = "".join(justificativas_banco)

# 2. Lógica da Governança
if foco_governanca == "Glossário, políticas e democratização (Foco em Negócios/Compliance)":
    camada_governanca = "Collibra" if modelo_licenca == "Apenas Comercial / Pago" else "OpenMetadata"
    justificativa_gov = f"A governança é liderada pelo {camada_governanca}, focando na democratização e no glossário de negócios. "
elif foco_governanca == "Qualidade de dados automatizada e Master Data Management (Foco em MDM)":
    camada_governanca = "Ataccama ONE" if modelo_licenca == "Apenas Comercial / Pago" else "Apache Atlas"
    justificativa_gov = f"A governança foca em confiabilidade e MDM através do {camada_governanca}, garantindo dados limpos. "
else: 
    # Foco em TI e Linhagem Técnica
    if modelo_licenca == "Apenas Comercial / Pago":
        camada_governanca = "Rocket Data Intelligence" 
    else:
        camada_governanca = "DataHub" if "Streaming" in cenario_dados else "OpenMetadata"
    justificativa_gov = f"O {camada_governanca} foi escolhido para rastrear a linhagem técnica ponta a ponta para a engenharia. "

# 3. Lógica do Ecossistema (Ingestão e Processamento)
if modelo_licenca == "Apenas Comercial / Pago":
    if "Streaming" in cenario_dados:
        camada_ingestao, camada_processamento = "Apache Kafka", "Apache Spark"
        justificativa_eco = "O Kafka gerencia a mensageria em tempo real, processado pelo Spark. "
    else: 
        camada_ingestao, camada_processamento = "Fivetran", "Databricks"
        justificativa_eco = "O Fivetran automatiza a ingestão de lotes, com transformações no Databricks. "
    camada_semantica_tool = "Looker" if exige_semantica else "Não implementada"
else:
    if "Streaming" in cenario_dados:
        camada_ingestao, camada_processamento = "Apache Kafka", "Apache Spark"
        justificativa_eco = "O ecossistema flui com Kafka na mensageria e Spark no processamento contínuo. "
    else: 
        camada_ingestao, camada_processamento = "Apache Airflow", "Trino + DuckDB"
        justificativa_eco = "O Airflow orquestra as rotinas, sendo processado pelo motor distribuído do Trino e DuckDB. "
    camada_semantica_tool = "Cube" if exige_semantica else "Não implementada"

# Junta a redação final dinamicamente
justificativa = justificativa_banco + justificativa_eco + justificativa_gov

st.write(f"**Visão Geral:** {justificativa}")


mermaid_code = f"""graph LR
    subgraph Ingestao ["Ingestão e Dados Brutos"]
        A["{camada_ingestao}"]
    end
    subgraph Armazenamento ["Armazenamento e Processamento"]
        B["{camada_armazenamento}"]
        C["{camada_processamento}"]
    end
    subgraph Consumo ["Consumo e Negócios"]
        D["{camada_semantica_tool}"]
        E["Aplicações (BI, Python, R)"]
    end
    subgraph Governanca ["Governança"]
        F["{camada_governanca}"]
    end
    A -->|"Carrega Dados"| B
    B -->|"Consulta"| C
"""
if exige_semantica:
    mermaid_code += "    C -->|\"Padroniza Métricas\"| D\n    D -->|\"Consome\"| E\n"
else:
    mermaid_code += "    C -->|\"Consome Direto\"| E\n"
mermaid_code += "    F -.- A\n    F -.- B\n    F -.- C\n    F -.- D\n    style F fill:#4a148c,stroke:#fff,color:#fff\n    style D fill:#00695c,stroke:#fff,color:#fff"

graphbytes = mermaid_code.encode("utf8")
base64_bytes = base64.b64encode(graphbytes)
base64_string = base64_bytes.decode('ascii')
st.image(f"https://mermaid.ink/svg/{base64_string}", use_container_width=True)

# --- 5. LÓGICA DE GERAÇÃO DO DOCUMENTO WORD (.DOCX) ---
def gerar_relatorio_word(df_filtrado, b64_diagrama):
    doc = Document()
    doc.add_heading('Relatório de Arquitetura de Dados', 0)
    
    # Metodologia
    doc.add_heading('1. Metodologia de Avaliação das Ferramentas', level=1)
    doc.add_paragraph("As notas de Suporte e Flexibilidade apresentadas neste relatório foram estabelecidas com base nos seguintes critérios técnicos do mercado:")
    doc.add_paragraph("• Suporte (1 a 5): Avalia o tempo de maturidade da ferramenta, o tamanho e o engajamento da comunidade (para soluções Open Source) ou o nível de serviço corporativo (SLA). Uma nota 5 indica que a ferramenta possui anos no mercado, vasta documentação oficial clara e uma comunidade ativa capaz de solucionar eventuais problemas rapidamente.")
    doc.add_paragraph("• Flexibilidade (1 a 5): Avalia o grau em que a ferramenta pode ser customizada e integrada ao ecossistema moderno. Uma nota 5 indica que ela oferece APIs abertas, integra-se de forma nativa com as principais linguagens de análise (Python, R, Java, SQL) e utiliza padrões abertos que evitam o aprisionamento tecnológico (vendor lock-in).")

    # Arquitetura Macro
    doc.add_heading('2. Visão Macro da Arquitetura', level=1)
    resumo_macro = (
        f"A figura abaixo apresenta o diagrama da arquitetura recomendada, passando pelas camadas de Governança, "
        f"Ingestão de dados brutos, Armazenamento e processamento, e Consumo e negócios. Para este projeto, "
        f"foi selecionado {camada_governanca} para a Governança; {camada_ingestao} para a Ingestão; "
        f"{camada_armazenamento} e {camada_processamento} para o Armazenamento e processamento; e "
        f"{camada_semantica_tool} para a padronização e consumo."
    )
    doc.add_paragraph(resumo_macro)
    doc.add_paragraph(f"Detalhe da topologia: {justificativa}")
    
    try:
        url_imagem = f"https://mermaid.ink/img/{b64_diagrama}?type=png"
        resposta = requests.get(url_imagem)
        if resposta.status_code == 200:
            doc.add_picture(BytesIO(resposta.content), width=Inches(6.0))
    except Exception:
        doc.add_paragraph("[Aviso: Imagem do diagrama indisponível offline.]")

    # Tabela com Links
    doc.add_heading('3. Tabela de Ferramentas Recomendadas', level=1)
    doc.add_paragraph("A tabela de Ferramentas Recomendadas abaixo apresenta todas as ferramentas que cumpriram os critérios de seleção do relatório (licenciamento e notas de corte configuradas), não apenas as ferramentas que foram indicadas no diagrama visual principal.")
    
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = 'Table Grid'
    hdr = tabela.rows[0].cells
    hdr[0].text = 'Ferramenta'
    hdr[1].text = 'Categoria'
    hdr[2].text = 'Resumo'
    hdr[3].text = 'Site / Link'
    
    for _, row in df_filtrado.iterrows():
        row_cells = tabela.add_row().cells
        row_cells[0].text = str(row['nome'])
        row_cells[1].text = str(row['categoria'])
        row_cells[2].text = str(row['resumo'])
        row_cells[3].text = str(row['link'])

    # Justificativas Dinâmicas das Ferramentas
    doc.add_heading('4. Justificativa Detalhada por Ferramenta', level=1)
    
    # Criar string unificada da arquitetura para descobrir quem é "titular"
    str_arquitetura = f"{camada_ingestao} {camada_armazenamento} {camada_processamento} {camada_semantica_tool} {camada_governanca}"
    
    # Mapear titulares por categoria
    titulares_cat = {}
    for _, r in df_filtrado.iterrows():
        if r['nome'] in str_arquitetura:
            c = r['categoria']
            if c not in titulares_cat:
                titulares_cat[c] = []
            titulares_cat[c].append(r['nome'])
            
    for _, row in df_filtrado.iterrows():
        doc.add_heading(f"{row['nome']} ({row['categoria']})", level=2)
        
        texto_sup = ""
        if row['suporte'] == 5: texto_sup = "é amplamente consolidada, com documentação farta, longo histórico de estabilidade no mercado e uma comunidade engajada para resolução imediata de problemas."
        elif row['suporte'] == 4: texto_sup = "possui ótima documentação e canais ativos que resolvem a imensa maioria dos casos de uso de pesquisa."
        elif row['suporte'] <= 3: texto_sup = "oferece o suporte básico necessário, porém pode exigir maior conhecimento técnico da equipe interna para configurações avançadas."
        
        texto_flex = ""
        if row['flexibilidade'] == 5: texto_flex = "é altamente customizável, integra-se com facilidade a múltiplas ferramentas analíticas, aceita linguagens variadas e foca em padrões abertos."
        elif row['flexibilidade'] == 4: texto_flex = "oferece excelentes conexões nativas com o ecossistema moderno de dados e atende à maioria dos padrões de engenharia."
        elif row['flexibilidade'] <= 3: texto_flex = "possui um ecossistema mais contido, operando de forma muito eficiente dentro do seu propósito específico, com menos opções de customização extrema."

        # Checa se a ferramenta está no diagrama (Titular) ou não (Alternativa)
        if row['nome'] in str_arquitetura:
            paragrafo = f"Foi escolhida a ferramenta {row['nome']} para compor o diagrama principal na arquitetura porque ela atende perfeitamente aos requisitos do projeto ({row['resumo']}). "
            paragrafo += f"Sua nota {row['suporte']} para suporte indica que ela {texto_sup} "
            paragrafo += f"Além disso, sua nota {row['flexibilidade']} para flexibilidade indica que ela {texto_flex}"
        else:
            tits = titulares_cat.get(row['categoria'], [])
            if tits:
                nomes_tits = " e ".join(tits)
                paragrafo = f"A ferramenta escolhida para a arquitetura principal foi a {nomes_tits}, mas a {row['nome']} pode substituir por também ser {row['tipo'].lower()} e especializada em {row['categoria']}. "
            else:
                paragrafo = f"Esta ferramenta não apareceu no diagrama principal (que priorizou o cenário de dados selecionado nos filtros), mas a {row['nome']} é uma excelente opção {row['tipo'].lower()} de {row['categoria']}. "
            
            paragrafo += f"Com uma nota {row['suporte']} para suporte e nota {row['flexibilidade']} para flexibilidade, ela é uma ótima alternativa caso não seja possível adotar a arquitetura padrão. Isso porque ela {texto_flex} E no quesito de suporte, ela {texto_sup}"
            
        doc.add_paragraph(paragrafo)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 6. TABELA DETALHADA E BOTÕES DE EXPORTAÇÃO ---
st.divider()
st.subheader("📊 Ecossistema e Exportação")

# Filtra o DataFrame dinamicamente
recomendacoes = df.copy()
if filtro_categoria != "Todas": recomendacoes = recomendacoes[recomendacoes['categoria'] == filtro_categoria]
if modelo_licenca == "Apenas Open Source": recomendacoes = recomendacoes[recomendacoes['tipo'] == "Open Source"]
elif modelo_licenca == "Apenas Comercial / Pago": recomendacoes = recomendacoes[recomendacoes['tipo'] == "Comercial / Pago"]
recomendacoes = recomendacoes[(recomendacoes['suporte'] >= nota_suporte) & (recomendacoes['flexibilidade'] >= nota_flexibilidade)]

if recomendacoes.empty:
    st.warning("Nenhuma ferramenta atende aos critérios selecionados.")
else:
    recomendacoes['Nota Final'] = recomendacoes['suporte'] + recomendacoes['flexibilidade'] + recomendacoes['facilidade_instalacao']
    recomendacoes = recomendacoes.sort_values(by='Nota Final', ascending=False)
    
    st.dataframe(
        recomendacoes[['nome', 'categoria', 'tipo', 'resumo', 'suporte', 'flexibilidade', 'link']],
        column_config={
            "nome": "Ferramenta", "categoria": "Categoria", "tipo": "Licença", "resumo": "O que faz?",
            "suporte": "Suporte (1-5)", "flexibilidade": "Customização (1-5)",
            "link": st.column_config.LinkColumn("Site Oficial")
        },
        hide_index=True, use_container_width=True
    )
    
    col1, col2 = st.columns([1, 1])
    with col1:
        csv_data = recomendacoes.to_csv(index=False, sep=';', encoding='utf-8-sig')
        st.download_button("📥 Baixar Tabela em CSV", data=csv_data, file_name="ferramentas.csv", mime="text/csv")
    with col2:
        relatorio_docx = gerar_relatorio_word(recomendacoes, base64_string)
        st.download_button("📄 Baixar Relatório Completo (.docx)", data=relatorio_docx, file_name="Relatorio_Arquitetura.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")